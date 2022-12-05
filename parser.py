import json
import os.path
import pickle
import time
import pandas as pd

import nltk
from docx import Document
from nltk.corpus.reader.api import CorpusReader, CategorizedCorpusReader
from nltk.corpus.reader.plaintext import CategorizedPlaintextCorpusReader, sent_tokenize
from nltk import wordpunct_tokenize, pos_tag
from readability.readability import Unparseable
from readability.readability import Document as Paper
import bs4

# import nltk

# nltk.download('punkt')
# nltk.download('averaged_perceptron_tagger')
# nltk.download('averaged_perceptron_tagger_ru')

CAT_PATTERN = r'([\w_\s\.]+)/.*'
DOC_PATTERN = r'(?!\.)[\w_\s]+/[\w\s\d\-]+\.txt'
# HTML_PATTERN = r'(?!\.)[/\w_\s]+[/\w+.-]*[\.html]+'
HTML_PATTERN = r'.*[\.html]+'
# PKL_PATTERN = r'(?!\.)[a-z_\s]+/[a-f0-9]+\.pickle'
PKL_PATTERN = r'.*[\.pickle]+'
TAGS = ['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'h7', 'p', 'li']


# TODO: Убрать дубликаты в коде, выделить в отдельные функции.
#  P.S. Аккуратно с наследованием!

def iscyrillic(s):
    cyrillic_symbols = set('абвгдеёжзийклмнопрстуфхцчшщъыьэюя-')
    word_symbols = set(s.lower())
    return word_symbols.difference(cyrillic_symbols) == set()


class TXTCorpusReader(CategorizedPlaintextCorpusReader):
    def __init__(self, root, fileids=DOC_PATTERN, encoding='utf8', **kwargs):
        """
        Инициализация класса для чтения TXT-файлов с использованием средств NLTK.

        :param root: Путь к файлу
        :param fileids: Шаблон по которому извлекаются документы
        :param encoding: Кодировка, в которой считываются файлы
        :param kwargs: Дополнительные параметры
        """
        if not any(key.startswith('cat_') for key in kwargs.keys()):
            kwargs['cat_pattern'] = CAT_PATTERN
        CategorizedPlaintextCorpusReader.__init__(self, root, fileids, encoding, **kwargs)
        # CorpusReader.__init__(self, root, fileids, encoding)

    def resolve(self, fileids=None, categories=None):
        """
        Возвращает список идентификаторов файлов или названий категорий,
        которые передаются каждой внутренней функции объекта чтения корпуса.
        """
        # if fileids is not None and categories is not None:
        #     raise ValueError("Укажите id-файлов или категории (выберите один вариант)")
        if categories is not None:
            return self.fileids(categories)
        return fileids

    def docs(self, fileids=None, categories=None):
        """
        Генератор, который возвращает весь текст документа закрывая его по завершению чтения.
        """
        fileids = self.resolve(fileids, categories)

        for path, encoding in self.abspaths(fileids, include_encoding=True):
            with open(path, 'r', encoding=encoding) as f:
                yield f.read()

    def sizes(self, fileids=None, categories=None):
        """
        Генератор, возвращающий имена и размеры файлов.
        """
        fileids = self.resolve(fileids, categories)

        for path in self.abspaths(fileids):
            yield path, os.path.getsize(path)

    def paras(self, fileids=None, categories=None):
        """
        Генератор для выделения абзацев из документа.
        """
        fileids = self.resolve(fileids, categories)
        for doc in self.docs(fileids, categories):
            lines = doc.split('\n')
            for line in lines:
                if line != '':
                    yield line

    def sents(self, fileids=None, categories=None):
        """
        Выделяет предложения из абзацев.
        """
        for paragraph in self.paras(fileids, categories):
            for sentence in sent_tokenize(paragraph):
                yield sentence

    def words(self, fileids=None, categories=None):
        """
        Выделяет слова из предложений.
        """
        for sentence in self.sents(fileids, categories):
            for token in wordpunct_tokenize(sentence):
                yield token

    def tokenize(self, fileids=None, categories=None):
        for paragraph in self.paras(fileids, categories):
            yield [pos_tag(wordpunct_tokenize(sent), lang='rus') for sent in sent_tokenize(paragraph)]


class HTMLCorpusReader(CategorizedCorpusReader, CorpusReader):
    """
    Объект для чтения корпуса с HTML-документами для получения возможности
    дополнительной предварительной обработки.
    """

    def __init__(self, root, fileids=HTML_PATTERN, encoding='utf8', tags=TAGS, **kwargs):
        """
        Инициализация класса для чтения HTML-файлов с использованием средств NLTK.

        :param root: Путь к файлу
        :param fileids: Шаблон по которому извлекаются документы
        :param encoding: Кодировка, в которой считываются файлы
        :param tags: HTML теги использующиеся для извлечения текста
        :param kwargs: Дополнительные параметры
        """
        if not any(key.startswith('cat_') for key in kwargs.keys()):
            kwargs['cat_pattern'] = CAT_PATTERN

        # TXTCorpusReader.__init__(self, root, fileids, encoding, **kwargs)
        CategorizedCorpusReader.__init__(self, kwargs)
        CorpusReader.__init__(self, root, fileids, encoding)
        self.tags = tags

    def resolve(self, fileids=None, categories=None):
        """
        Возвращает список идентификаторов файлов или названий категорий,
        которые передаются каждой внутренней функции объекта чтения корпуса.
        """
        if fileids is not None and categories is not None:
            raise ValueError("Укажите id-файлов или категории (выберите один вариант)")
        if categories is not None:
            return self.fileids(categories)
        if fileids is not None:
            return fileids
        return self.fileids()

    def docs(self, fileids=None, categories=None):
        """
        Генератор, который возвращает весь текст документа закрывая его по завершению чтения.
        """
        fileids = self.resolve(fileids, categories)

        for path, encoding in self.abspaths(fileids, include_encoding=True):
            with open(path, 'r', encoding=encoding) as f:
                yield f.read()

    def html(self, fileids=None, categories=None):
        """
        Возвращает содержимое HTML документов, очищая их с помощью библиотеки readability-lxml.
        """
        for doc in self.docs(fileids, categories):
            try:
                yield Paper(doc).summary()
            except Unparseable as e:
                print(f'Не могу распарсить HTML-страницу: {e}')
                continue

    def paras(self, fileids=None, categories=None):
        """
        Генератор для выделения абзацев из HTML.
        Использует библиотеку BeautifulSoup.
        """
        tags = ['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'h7', 'p', 'li']
        fileids = self.resolve(fileids, categories)
        for html in self.html(fileids, categories):
            soup = bs4.BeautifulSoup(html, 'lxml')
            for element in soup.find_all(tags):
                yield element.text
            soup.decompose()

    def sents(self, fileids=None, categories=None):
        """
        Выделяет предложения из абзацев.
        """
        for paragraph in self.paras(fileids, categories):
            for sentence in sent_tokenize(paragraph):
                yield sentence

    def words(self, fileids=None, categories=None):
        """
        Выделяет слова из предложений.
        """
        for sentence in self.sents(fileids, categories):
            for token in wordpunct_tokenize(sentence):
                yield token

    def tokenize(self, fileids=None, categories=None):
        for paragraph in self.paras(fileids, categories):
            yield [pos_tag(wordpunct_tokenize(sent), lang='rus') for sent in sent_tokenize(paragraph)]

    # TODO: (*дубликаты*)
    # def token_filter(self, tokens):
    #     return filtered_tokens

    def describe(self, fileids=None, categories=None):
        """
        Выполняет обход содержимого корпуса и возвращает словарь с различными оценками,
        описывающим состояние корпуса.

        :param (str) fileids: Название отдельно выбранного файла
        :param (str) categories: Категория, из которой извлекаются все документы
        :return (dict): Словарь с различными оценками,
            описывающим состояние корпуса.
        """
        started = time.time()

        counts = nltk.FreqDist()
        tokens = nltk.FreqDist()

        for para in self.paras(fileids, categories):
            counts['paras'] += 1

            # for sent in t_para:
            for sent in sent_tokenize(para):
                counts['sents'] += 1

                # Удаление символов
                delete_symbols = ['\n', ',', '.', '!', '?', ':', ';', '_', '#', '(', ')', '\\', '/', '⇡', '×',
                                  '«', '»', '', '—', '<', '>', '+', '[', ']', '👎', '❌', '🇵🇹', '🇦🇷', '🟢', '🔴',
                                  '⚪️', '\u200b', '…', '👌', '•\u200e', '"', '\u2009']
                word_tokens = str(sent)
                word_tokens = word_tokens.replace('\xa0', ' ')
                for ds in delete_symbols:
                    word_tokens = word_tokens.replace(ds, '')
                word_tokens = word_tokens.split(' ')

                # for word in wordpunct_tokenize(sent):
                for word in word_tokens:
                    counts['words'] += 1
                    tokens[word] += 1

        n_fileids = len(self.resolve(fileids, categories) or self.fileids)
        n_topics = len(self.categories(self.resolve(fileids, categories)))

        result = {
            'files': n_fileids,
            'topics': n_topics,
            'paras': counts['paras'],
            'sents': counts['sents'],
            'words': counts['words'],
            'vocab': len(tokens),
            'lexdiv': float(counts['words']) / float(len(tokens)),
            'ppdoc': float(counts['paras']) / float(n_fileids),
            'sspar': float(counts['sents']) / float(counts['paras']),
            'secs': time.time() - started
        }
        return result


class Preprocessor(object):
    """
    Обертывает 'HTMLCorpusReader' и выполняет лексемизацию
    с маркировкой частями речи.
    """

    def __init__(self, corpus, target='./processed_corpus', **kwargs):
        self.corpus = corpus
        self.target = target

    def fileids(self, fileids=None, categories=None):
        fileids = self.corpus.resolve(fileids, categories)
        if fileids:
            return fileids
        return self.corpus.fileids()

    def abspath(self, fileid):
        """
        Возвращает путь к файлу относительно корня целевого корпуса.
        """
        # Найти путь к каталогу относительно корня искходного корпуса.
        parent = os.path.relpath(os.path.dirname(self.corpus.abspath(fileid)), self.corpus.root)

        # Выделить части пути для реконструирования
        basename = os.path.basename(fileid)
        name, ext = os.path.splitext(basename)

        # Сконструировать имя файла с расширением .pickle
        basename = name + '.pickle'

        # Вернуть путь к файлу относительно корня целевого корпуса
        return os.path.normpath(os.path.join(self.target, parent, basename))

    def tokenize(self, fileid):
        """
        Обработка и токенизация слов из указанного файла.
        """
        # TODO: переделать для русского языка (корректно)
        for paragraph in self.corpus.paras(fileids=fileid):
            yield [pos_tag(wordpunct_tokenize(sent), lang='rus') for sent in sent_tokenize(paragraph)]

    def process(self, fileid):
        """
        Вызывается для одного файла, проверяет местоположение на диске, чтобы
        гарантировать отсутствие ошибок, использует +tokenize()+ для предварительной
        обработки и записывает трансформированный документ в виде сжатого архива.
        """
        # Определить путь к файлу для записи результата
        target = self.abspath(fileid)
        parent = os.path.dirname(target)

        # Убедиться в сущестовании каталога
        if not os.path.exists(parent):
            os.makedirs(parent)

        # Убедиться, что parent - это каталог, а не файл
        if not os.path.isdir(parent):
            raise ValueError('Пожалуйста, убедитесь что указан каталог, а не файл')

        # Создать структуру данных для записи в архив
        doc = self.tokenize(fileid)
        document = list(doc)

        # Записать данные в архив на диск
        with open(target, 'wb') as file:
            pickle.dump(document, file, pickle.HIGHEST_PROTOCOL)

        # Удалить документ из памяти
        del document

        # Вернуть путь к целевому файлу
        return target

    def transform(self, fileids=None, categories=None):
        # Создать целевой каталог, если чего ещё нет
        if not os.path.exists(self.target):
            os.makedirs(self.target)

        # Получить имена файлов для обработки
        for fileid in self.fileids(fileids, categories):
            yield self.process(fileid)

    def save_tokens(self, fileids=None, categories=None):
        """
        Сохранение простого словаря со словами в файле формата JSON.
        """
        tokens = nltk.FreqDist()
        counts = 0
        for para in self.corpus.paras(fileids, categories):
            for sent in sent_tokenize(para):
                # Удаление символов
                delete_symbols = ['\n', ',', '.', '!', '?', ':', ';', '_', '#', '(', ')', '\\', '/', '⇡', '×',
                                  '«', '»', '', '—', '<', '>', '+', '%', '[', ']', '👎', '❌', '🇵🇹', '🇦🇷', '🟢',
                                  '🔴', '⚪️', '\u200b', '…', '👌', '•\u200e', '"', '\u2009', '£']
                word_tokens = str(sent)
                word_tokens = word_tokens.replace('\xa0', ' ')
                for ds in delete_symbols:
                    word_tokens = word_tokens.replace(ds, '')
                word_tokens = word_tokens.split(' ')
                counts += 1
                # print(counts)
                for word in word_tokens:
                    if iscyrillic(word) and word != '':
                        tokens[word] += 1

        word_list = {
            "words": list(tokens.keys())
        }

        path = f'{self.target}/word_list_ru_050922.json'

        with open(path, 'w', encoding='utf-8') as output_file:
            json.dump(word_list, output_file)

        return path


# TODO: описание функций для извлечения данных
class PickledCorpusReader(HTMLCorpusReader):
    def __init__(self, root, fileids=PKL_PATTERN, **kwargs):
        if not any(key.startswith('cat_') for key in kwargs.keys()):
            kwargs['cat_pattern'] = CAT_PATTERN
        CategorizedCorpusReader.__init__(self, kwargs)
        CorpusReader.__init__(self, root, fileids)

    def docs(self, fileids=None, categories=None):
        fileids = self.resolve(fileids, categories)

        for path in self.abspaths(fileids):
            with open(path, 'rb') as file:
                yield pickle.load(file)

    def paras(self, fileids=None, categories=None):
        for doc in self.docs(fileids, categories):
            for para in doc:
                yield para

    def sents(self, fileids=None, categories=None):
        for para in self.paras(fileids, categories):
            for sent in para:
                yield sent

    def tagged(self, fileids=None, categories=None):
        for sent in self.sents(fileids, categories):
            for tagged_token in sent:
                yield tagged_token

    def words(self, fileids=None, categories=None):
        for tagged in self.tagged(fileids, categories):
            yield tagged[0]


def create_docx(path_to_docx, text_structure, folder_file_name):
    """
    Функция для обработки .pickle и с последующим созданием файлов .docx с текстовой информацией
    :param path_to_docx:
    :param text_structure:
    :param folder_file_name:
    :return:
    """
    name_list = folder_file_name.split('/')
    folder = name_list[0]
    file = name_list[1].replace('.pickle', '')
    document = Document()
    for ts in text_structure:
        if ts is None:
            document.add_paragraph()
        else:
            for para in ts:
                if para:
                    line = ''
                    for word in para:
                        line += str(f' {word[0]}')
                    document.add_paragraph(line)
    if not os.path.exists(f'{path_to_docx}/{folder}'):
        os.makedirs(f'{path_to_docx}/{folder}')

    if not os.path.isdir(f'{path_to_docx}/{folder}'):
        raise ValueError('Пожалуйста, убедитесь что указан каталог, а не файл')

    document.save(f'{path_to_docx}/{folder}/{file}.docx')
    return 0


def resulted_xlsx(path_to_xlsx, fileids, pickled_data):
    """
    Функция для обработки .pickle и с последующим созданием таблицы с классами в формате .xlsx
    :param path_to_xlsx:
    :param fileids:
    :param pickled_data:
    """
    resulted_list = []
    site_dict = {
        '3dnews.ru': '3dnews.ru',
        'kuban.rbc.ru': 'kuban.rbc.ru',
        'ria.ru': 'ria.ru',
        'lenta.rupartsnews': 'lenta.ru/parts/news',
        'sports.ru': 'sports.ru',
    }
    filter_list = ['https', '/', 'html', 'meta', 'head', 'xn', 'ria', 'ru', 'internet']

    df = pd.DataFrame(resulted_list, columns=['Содержание', 'Источник', 'Да', '?', 'Нет'])
    df.to_excel(f'{path_to_xlsx}/resulted_table.xlsx', index=False)

    # if not os.path.exists(f'{path_to_docx}/{folder}'):
    #     os.makedirs(f'{path_to_docx}/{folder}')
    #
    # if not os.path.isdir(f'{path_to_docx}/{folder}'):
    #     raise ValueError('Пожалуйста, убедитесь что указан каталог, а не файл')
    #
    # document.save(f'{path_to_docx}/{folder}/{file}.docx')
    return 0


if __name__ == "__main__":
    # TODO: Разбить по функциям
    html_reader = HTMLCorpusReader('./sites', fileids=HTML_PATTERN, cat_pattern=CAT_PATTERN, tags=TAGS)
    # category_type = 'it'
    # fileid = f'./{category_type}/Alphacool_Eisblock_Acryl_GPX_Zotac_RTX_3070_Ti.html'
    # print(html_reader.describe())

    # preproc = Preprocessor(html_reader)
    # started = time.time()
    # preproc.save_tokens()
    # print(time.time() - started)

    # Сохранение данных
    # preproc = Preprocessor(html_reader)
    # fileids = preproc.fileids()
    # print('--------------------- HTML to PICKLE ---------------------')
    # for i, fileid in enumerate(fileids):
    #     print(f'{i:6}: {fileid}')
    #     preproc.process(fileid)
    # pass

    # Чтение данных из архива
    # fileid = f'./{category_type}/Alphacool_Eisblock_Acryl_GPX_Zotac_RTX_3070_Ti.pickle'
    # fileid = f'./3dnews.ru/3dnews.ru1059046obzor-huawei-p50-propage-1.html_1744_09032022.pickle'
    pick = PickledCorpusReader('./processed_corpus', fileids=PKL_PATTERN, cat_pattern=CAT_PATTERN, tags=TAGS)
    fileids = pick.resolve()
    # path_to_docx = './docx_result'
    # print('--------------------- PICKLE to DOCX ---------------------')
    # for i, fileid in enumerate(fileids):
    #     print(f'{i:6}: {fileid}')
    #     text_structure = next(pick.docs(fileid))
    #     create_docx(path_to_docx=path_to_docx, text_structure=text_structure, folder_file_name=fileid)
    # pass

    print('--------------------- PICKLE to XLSX ---------------------')
    path_to_xlsx = './docx_result'
    resulted_xlsx(path_to_xlsx=path_to_xlsx, fileids=fileids, pickled_data=pick)
    pass

    # Сохранение словаря с уникальными словами
    # path = f'./processed_corpus/word_list_ru.json'
    # with open(path) as f:
    #     template = json.load(f)
    # print(template)
    # print(len(template['words']))
    # pass

# ----------------------------------------------------------------------------------------------------------------------
# TXT-Reader test
# ----------------------------------------------------------------------------------------------------------------------
# corpus = CategorizedPlaintextCorpusReader('./corpus', DOC_PATTERN, cat_pattern=CAT_PATTERN)
#
# print(corpus.categories())
# print(corpus.fileids())

# txt_reader = TXTCorpusReader('./corpus', fileids=DOC_PATTERN, cat_pattern=CAT_PATTERN)
# # print(txt_reader.resolve())
# category_type = 'auto'
# # print(next(txt_reader.tokenize(fileids=f'./{category_type}/audi_rs4_5.txt')))
# for trt in txt_reader.tokenize(fileids=f'./{category_type}/audi_rs4_5.txt'):
#     print(trt)
# print(next(txt_reader.tokenize(fileids=f'./{category_type}/audi_rs4_5.txt')))
# for tr in txt_reader.words(categories='auto'):
# for tr in txt_reader.sents(fileids='./it/news_nvidia_amd_1.txt'):
# for tr in txt_reader.paras():
#     print(tr)
# print(next(txt_reader.paras()))
# for t in txt_reader.docs():
#     print(t)

# ----------------------------------------------------------------------------------------------------------------------
# HTML-Reader test
# ----------------------------------------------------------------------------------------------------------------------

# html_reader = HTMLCorpusReader('./corpus', fileids=DOC_PATTERN, cat_pattern=CAT_PATTERN, tags=TAGS)
#
# category_type = 'it'
# print(html_reader.describe(fileids=f'./{category_type}/Alphacool_Eisblock_Acryl_GPX_Zotac_RTX_3070_Ti.html'))
